import pandas as pd
import json
import tempfile
import os
import shutil
import zipfile
import google.generativeai as genai
import time
import docx
from .gemini_core import GeminiCore

class RiskMapper:
    def __init__(self, api_key):
        self.client = GeminiCore(api_key)

    def extract_text_from_docx(self, docx_path):
        try:
            doc = docx.Document(docx_path)
            full_text = []
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                     full_text.append(para.text)
            
            # Extract tables (naive text extraction)
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    full_text.append(" | ".join(row_data))
                    
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error reading DOCX {docx_path}: {e}")
            return ""

    def process(self, target_docx, target_csv, mapping_files, ref_files=[], logger=print):
        try:
            # Setup Request Directory
            PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            RESULT_DIR = os.path.join(PROJECT_ROOT, "data", "result")
            if not os.path.exists(RESULT_DIR):
                os.makedirs(RESULT_DIR)

            # Load Target CSV (to get Risk Names)
            try:
                df_target = pd.read_csv(target_csv)
            except:
                df_target = pd.read_excel(target_csv) # Fallback
            
            # Prioritize standard column names (Exact Match)
            if '위험률명' in df_target.columns:
                risk_col = '위험률명'
            elif 'Risk Name' in df_target.columns:
                risk_col = 'Risk Name'
            else:
                # Fuzzy Search
                possible_cols = [c for c in df_target.columns if '위험률명' in c or 'Risk Name' in c]
                if not possible_cols:
                     possible_cols = [c for c in df_target.columns if '위험률' in c or 'Risk' in c]
                risk_col = possible_cols[0] if possible_cols else df_target.columns[0]
            
            logger(f"Identified 'Risk Name' column: {risk_col}")
            
            # Keep original relevant columns (e.g., Risk ID)
            # We will merge results back to df_target based on risk_col
            target_risks = df_target[risk_col].dropna().unique().tolist()
            logger(f"Target Risks to process: {len(target_risks)} items.")

            # Prepare Mapping Context
            mapping_context = ""
            uploaded_maps = []
            if not isinstance(mapping_files, list): mapping_files = [mapping_files]

            for mf in mapping_files:
                try:
                    df_map = pd.read_csv(mf) if mf.endswith('.csv') else pd.read_excel(mf)
                    csv_text = df_map.to_csv(index=False)
                    mapping_context += f"\n\n=== Mapping Table: {os.path.basename(mf)} ===\n{csv_text}\n"
                    
                    # Create temp file for upload if needed (or just verify read)
                    # We will upload mapping files to Gemini as well for better context handling
                    # But if they are small, text context is fine. 
                    # Let's upload them to be safe if Gemini supports it well, or use text context if small.
                    # Given previous context, text context in prompt often works well for logic extraction.
                    # But for large tables, file upload is better.
                    # Let's stick to text context for prompt for now unless huge.
                except Exception as e:
                    logger(f"Error reading mapping file {mf}: {e}")

            # --- Phase 1: Logic Inference ---
            logic_result = "No Logic Extracted"
            
            # Group reference files
            ref_pairs = self._group_files_by_pair(ref_files)
            extracted_logics = []

            if ref_pairs:
                logger("\n=== Phase 1: Logic Inference ===")
                for base_name, files in ref_pairs.items():
                    # Identify which is docx and which is csv
                    ref_docx = next((f for f in files if f.endswith('.docx') or f.endswith('.doc')), None)
                    ref_csv = next((f for f in files if f.endswith('.csv') or f.endswith('.xlsx')), None)

                    if not ref_docx or not ref_csv:
                        logger(f"Skipping pair {base_name}: Missing DOCX or CSV/Excel.")
                        continue
                    
                    logger(f"Analyzing Pair: {base_name}...")
                    
                    # Read Reference Data
                    ref_doc_text = self.extract_text_from_docx(ref_docx)
                    try:
                        ref_df = pd.read_csv(ref_csv) if ref_csv.endswith('.csv') else pd.read_excel(ref_csv)
                        ref_csv_text = ref_df.to_csv(index=False)
                    except:
                        ref_csv_text = "Error reading CSV"

                    # Prompt for Logic Extraction
                    prompt_phase1 = f"""
                    **Role**: 보험 상품 설계 및 약관 분석 전문가 (Insurance Product Design & Policy Analysis Expert)
                    
                    **Goal**: 
                    '약관(Reference Policy)'의 내용과 그에 대응하는 '보장기준정보(Reference Coverage Data)' 간의 관계를 분석하여,
                    **'위험률명(Risk Name)'**을 기준으로 **'보장기준정보(Item Categories & Values)'**를 생성하는 **상세 추론 로직(Inference Logic)**을 도출하세요.

                    **Core Task**:
                    약관의 문구로부터 위험률별로 어떤 '항목구분(Item Category)'들이 생성되어야 하고, 각 항목의 값은 어떻게 매핑되는지 규칙을 찾아내야 합니다.

                    **Inputs**:
                    1. **Reference Policy (Text)**: 약관 내용 (일부 발췌)
                    2. **Reference Coverage Data (CSV)**: 해당 약관에 대한 정답 데이터 (위험률명, 항목구분상세, 항목값상세 등 포함)
                    3. **Mapping Definitions (Code Master)**: 코드 매핑 정의 (항목구분코드.csv, 그룹코드.csv, 추가구분키.csv 등 포함)
                    {mapping_context}
                    
                    **Reference Policy Content**:
                    {ref_doc_text}
                    
                    **Reference Coverage Data**:
                    {ref_csv_text}
                    
                    **Analysis Requirements**:
                    1. **Row Generation Logic (Variable Structure)**: 
                       - 위험률명(Risk Name) 별로 생성되는 행(Row)의 개수와 종류(항목구분명)가 다릅니다.
                       - 예: '뇌질환입원율' -> {{"보장유형", "지급유형", "한도유형", "부책진단그룹", "면책이벤트그룹"}} (5개 행)
                       - 예: '독감입원율' -> {{"보장유형", "지급유형", "한도유형", "부책진단그룹"}} (4개 행 - 면책이벤트그룹 없음)
                       - **어떤 기준(약관의 문구, 담보의 성격 등)으로 특정 항목구분(예: 면책이벤트그룹)이 추가되거나 빠지는지 원인을 분석하세요.**

                    2. **Value Inference & Mapping Logic**:
                       - 약관의 문구(Text)에서 '항목구분명'에 해당하는 '항목값명'을 어떻게 추출하는지 설명하세요.
                       - 추출된 '항목값명'을 **Code Master(매핑 테이블)**를 사용하여 '항목구분코드' 및 '항목값코드'로 변환하는 과정을 설명하세요.
                       - 보조 테이블('항목구분코드.csv', '그룹코드.csv', '추가구분키.csv')이 있다면 이를 어떻게 조합하여 코드를 찾는지 분석하세요.

                    **Output Format (Logic Description)**:
                    1. **[Condition for Item Category Generation]**: 위험률의 특징에 따라 어떤 '항목구분명'들이 생성되어야 하는지 조건 명시. (예: "면책 조항에 '특정 이벤트' 관련 내용이 있으면 '면책이벤트그룹' 항목을 추가한다.")
                    2. **[Mapping Rule per Category]**: 각 '항목구분명' 별로 값을 추출하고 코드를 매핑하는 구체적인 방법.
                    3. **[Keywords & Patterns]**: 약관에서 핵심적으로 봐야할 키워드 및 문장 패턴.
                    """

                    try:
                        resp = self.client.generate_content([prompt_phase1], request_options={'timeout': 600})
                        extracted_logics.append(f"=== Logic Derived from {base_name} ===\n{resp.text}")
                        logger(f"  > Logic extracted from {base_name}")
                    except Exception as e:
                        logger(f"  > Logic extraction failed for {base_name}: {e}")

            final_logic = "\n\n".join(extracted_logics)
            if not final_logic:
                final_logic = "기본 로직: 약관의 목차와 표를 참고하여 보장 항목을 식별하고, 매핑 테이블과 대조하여 코드를 부여한다."
            
            # Save Logic
            logic_file_path = os.path.join(RESULT_DIR, "Logic_Risk_Inference.txt")
            with open(logic_file_path, "w", encoding="utf-8") as f:
                f.write(final_logic)
            logger(f"Phase 1 Complete. Logic saved to {logic_file_path}")

            # --- Phase 2: Value Appliction ---
            logger("\n=== Phase 2: Value Inference ===")
            
            # Upload Target Policy (Docx) - Use text extraction or file upload if supported
            # Gemini supports PDF well, Docx support is via text extraction usually or converted to PDF.
            # Client.upload_file supports specific MIME types. DOCX is not always directly supported for analysis like PDF.
            # Safe bet: Extract text and pass as context, or chunk if too big.
            # Or convert DOCX to PDF?
            # Let's rely on text extraction for now as we did for reference.
            target_doc_text = self.extract_text_from_docx(target_docx)
            
            results = []
            
            # Process in batches of Risk Names to avoid context overflow
            batch_size = 5
            for i in range(0, len(target_risks), batch_size):
                batch_risks = target_risks[i:i+batch_size]
                logger(f"Processing Batch {i//batch_size + 1}: {batch_risks}")
                
                prompt_phase2 = f"""
                **Role**: 보험 약관 분석 및 보장기준정보 생성 AI Agent
                
                **Goal**: 
                제공된 **'Target Policy'**와 **'Inference Logic'**을 바탕으로, **'Target Risks'** 각각에 대한 상세 보장기준정보(Rows)를 생성하세요.
                
                **Logic (Extraction Rules)**:
                {final_logic}
                
                **Mapping Tables (Code Master)**:
                {mapping_context}
                
                **Target Policy Content**:
                {target_doc_text}
                
                **Target Risks (위험률명 리스트)**:
                {batch_risks}
                
                **Instructions**:
                1. **Search & Analyze**: 각 '위험률명(Risk Name)'에 해당하는 약관 내용을 정밀하게 찾아 분석하세요.
                2. **Apply Logic (Variable Rows)**: Phase 1에서 도출된 로직에 따라, 해당 위험률명에 필요한 모든 **'항목구분명'**을 결정하고 행(Row)을 생성하세요. 
                   - (중요: 위험률마다 생성되는 행의 개수가 다를 수 있습니다. 필요한 만큼 생성하세요.)
                3. **Extract & Map**: 생성된 각 '항목구분명'에 대해, 약관에서 '항목값명'을 결정하고, 매핑 테이블을 사용하여 **'항목구분코드', '항목값코드'**를 정확히 찾으세요.
                4. **Specific Rule (Coverage Identity & Anti-Patterns)**: 
                   - **Rule**: '보장유형'이 '특정치료비' 관련 값(예: '뇌심특정치료비...')이어야 하는 경우는, **'위험률명'에 반드시 '특정치료'라는 정확한 문구가 포함된 경우**로 한정합니다.
                   - **Anti-Pattern (Caution)**: 단순히 '치료', '혈전용해', '수술' 등의 단어가 겹친다고 해서 '특정치료비'로 매핑하면 안 됩니다.
                   - **Examples**:
                     - (O) '뇌혈관...**특정치료**Ⅱ발생률' -> '뇌심특정치료비(수술,혈전제거,혈전용해각연1회)' ('특정치료' 포함 O)
                     - (X) '급성심근경색증종합병원**혈전용해치료**발생률(연간1회한)' -> '질병치료' ('특정치료' 미포함, 단순 치료 행위임)
                5. **Reasoning**: 이 항목구분과 값이 선택된 구체적인 이유(약관 조항 인용 등)를 '선정근거' 항목에 적으세요.
                6. **Format**: 결과는 반드시 **JSON List** 형식이어야 합니다.
                7. **CRITICAL**: 결과 JSON의 "**위험률명**" 값은 반드시 입력으로 주어진 **'Target Risks'** 리스트에 있는 문자열과 **100% 동일(띄어쓰기 포함)**해야 합니다. 임의로 수정하거나 생략하지 마세요. (예: "암 입원일당" -> "암입원일당" (X))

                **Output JSON Structure**:
                [
                    {{
                        "위험률명": "Target Risk Name 1",
                        "항목구분코드": "Category Code 1",
                        "항목구분명": "Category Name 1 (e.g., 보장유형)",
                        "항목값코드": "Value Code 1",   
                        "항목값명": "Value Name 1 (e.g., 정액)",
                        "선정근거": "약관 제 3조 2항 '...지급한다'에 근거하여 정액 지급으로 판단."
                    }},
                    {{
                        "위험률명": "Target Risk Name 1",
                        "항목구분코드": "Category Code 2",
                        "항목구분명": "Category Name 2 (e.g., 부책진단그룹)",
                        "항목값코드": "Value Code 2",
                        "항목값명": "Value Name 2",
                        "선정근거": "별표 3 관련 질병 분류표 참조."
                    }},
                    ...
                    {{
                        "위험률명": "Target Risk Name 2",
                        ...
                    }}
                ]
                """
                
                try:
                    resp = self.client.generate_content([prompt_phase2], request_options={'timeout': 1000})
                    raw = resp.text.replace("```json", "").replace("```", "").strip()
                    if "[" in raw and "]" in raw:
                        json_str = raw[raw.find("["):raw.rfind("]")+1]
                        batch_data = json.loads(json_str)
                        results.extend(batch_data)
                        logger(f"  > Batch processed: {len(batch_data)} rows generated.")
                    else:
                        logger(f"  > Warning: Invalid JSON response for batch {batch_risks}")
                except Exception as e:
                    logger(f"  > Error processing batch {batch_risks}: {e}")
                
                time.sleep(2)

            # Save Result
            df_result = pd.DataFrame(results)
            
            # Normalization for Join
            # Ensure the risk column in result matches target/logic
            # Prompt output key was "위험률명", but we should Map it back to original risk_col
            if "위험률명" in df_result.columns and risk_col != "위험률명":
                df_result.rename(columns={"위험률명": risk_col}, inplace=True)

            # Ensure columns exist
            required_cols = [risk_col, "항목구분코드", "항목구분명", "항목값코드", "항목값명"]
            for c in required_cols:
                if c not in df_result.columns: df_result[c] = ""
            
            # Merge with original Target DataFrame to keep Risk ID and other info
            # Warning: df_result has multiple rows per risk. df_target has 1 row per risk.
            # Left Join: df_result (Many) <- df_target (One) ? No, we want key info from target.
            
            # Strategy: Merge df_result with df_target (subset)
            # df_target might have duplicates? Assuming unique risks in target list input usually.
            # If df_target has multiple rows for same risk, it's tricky. 
            # But usually it's unique list of risks to process.
            
            # Remove duplicates from right side (target info) just in case, on the join key
            df_target_unique = df_target.drop_duplicates(subset=[risk_col])
            
            # Merge
            # df_final = df_result left join df_target
            df_final = pd.merge(df_result, df_target_unique, on=risk_col, how='left')
            
            # Reorder columns: Risk ID/Name first, then calculated items
            # Identify ID column if exists (Explicit > Search)
            id_col = None
            if '위험률ID' in df_final.columns:
                id_col = '위험률ID'
            elif 'Risk ID' in df_final.columns:
                id_col = 'Risk ID'
            else:
                id_col = next((c for c in df_final.columns if 'ID' in c or '코드' in c and c != '항목구분코드' and c != '항목값코드'), None)
            
            final_cols = []
            if id_col and id_col != risk_col: final_cols.append(id_col)
            final_cols.append(risk_col)
            final_cols.extend(["항목구분코드", "항목구분명", "항목값코드", "항목값명"])
            
            # Add '선정근거' if present
            if "선정근거" in df_final.columns:
                final_cols.append("선정근거")
            
            # Check for any other metadata columns to keep? (e.g. description)
            # For now, stick to ID + Name + Results
            
            # Ensure final_cols exist
            final_cols = [c for c in final_cols if c in df_final.columns]
            
            df_final = df_final[final_cols]
            
            res_csv_path = os.path.join(RESULT_DIR, "Result_Risk_Inference.csv")
            df_final.to_csv(res_csv_path, index=False, encoding='utf-8-sig')
            logger(f"Phase 2 Complete. Result saved to {res_csv_path}")

            # Zip Results
            zip_path = os.path.join(RESULT_DIR, "PoC_Step4_Results.zip")
            with zipfile.ZipFile(zip_path, 'w') as zf:
                zf.write(logic_file_path, arcname="Logic_Risk_Inference.txt")
                zf.write(res_csv_path, arcname="Result_Risk_Inference.csv")
            
            return {
                "file_path": zip_path,
                "logic_path": logic_file_path,
                "preview": results[:5]
            }

        except Exception as e:
            logger(f"Critical Error: {e}")
            import traceback
            logger(traceback.format_exc())
            return {"error": str(e)}

    def _group_files_by_pair(self, file_paths):
        # Naive pairing by filename similarity or manual convention
        # As per request: "각 약관(워드) 파일들과 보장기준코드(csv) 파일들이 pairs 로 주어지며"
        # We assume filenames match or are clearly pairable.
        # Implementation: Group by filename stem (excluding extension)
        groups = {}
        for f in file_paths:
            base = os.path.splitext(os.path.basename(f))[0]
            # Handle potential suffixes like "_약관", "_코드" if strictly named
            # For now, exact match on stem
            if base not in groups: groups[base] = []
            groups[base].append(f)
        
        # Filter for valid pairs (at least 2 files)
        return {k: v for k, v in groups.items() if len(v) >= 2}
